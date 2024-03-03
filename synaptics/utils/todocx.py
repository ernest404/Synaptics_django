import docx
from docx.shared import RGBColor
from datetime import date
import io
import re

def clean_string(input_string):
    # Regular expression to match control characters and NULL bytes
    control_chars = re.compile(r'[\x00-\x1f\x7f-\x9f]')
    # Substitute matched characters with an empty string
    cleaned_string = re.sub(control_chars, '', input_string)
    return cleaned_string

def doc_builder(summary_list):

    # Create an instance of a word document
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading('MEDIA COVERAGE ON SECURITY ISSUES ', 0)

    doc.add_heading('Date: {}'.format(date.today().strftime("%d/%m/%Y")), 1).alignment = 2


    for i, summary in enumerate(summary_list):

        # Add paragraph
        doc.add_heading('News Report {}:'.format(i+1), 3)
        para = doc.add_paragraph().add_run(clean_string(summary))


    # Save the document to a location
    doc.save('downloads/summary report.docx')

def doc_builder(summary_list):
    doc = docx.Document()
    doc.add_heading('MEDIA COVERAGE ON SECURITY ISSUES', 0)
    doc.add_heading('Date: {}'.format(date.today().strftime("%d/%m/%Y")), 1).alignment = 2

    for i, summary in enumerate(summary_list):
        doc.add_heading('Summary {}:'.format(i + 1), 3)
        para = doc.add_paragraph().add_run(summary)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # Rewind the buffer
    return file_stream    